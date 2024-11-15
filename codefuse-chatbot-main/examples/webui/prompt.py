import streamlit as st
import os
import time
from datetime import datetime
import traceback
from typing import Literal, Dict, Tuple
from st_aggrid import AgGrid, JsCode
from st_aggrid.grid_options_builder import GridOptionsBuilder
import pandas as pd

# 新加
from io import BytesIO
from docx import Document


from .utils import *
from muagent.utils.path_utils import *
from muagent.service.service_factory import get_kb_details, get_kb_doc_details
from muagent.orm import table_init

def prompt_page(api: ApiRequest):
    # 判断表是否存在并进行初始化
    table_init()

    now = datetime.now()
    with st.sidebar:

        cols = st.columns(2)
        export_btn = cols[0]
        if cols[1].button(
                "清空prompt",
                use_container_width=True,
        ):
            st.experimental_rerun()

    export_btn.download_button(
        "导出记录",
        "测试prompt",
        file_name=f"{now:%Y-%m-%d %H.%M}_对话记录.md",
        mime="text/markdown",
        use_container_width=True,
    )


# def prompt_page(api: ApiRequest):
#     # 判断表是否存在并进行初始化
#     table_init()
#
#     now = datetime.now()
#     with st.sidebar:
#
#         cols = st.columns(2)
#         export_btn = cols[0]
#         if cols[1].button(
#                 "清空prompt",
#                 use_container_width=True,
#         ):
#             st.experimental_rerun()
#
#     # 创建 Word 文档
#     document = Document()
#     document.add_heading('对话记录', level=1)
#     document.add_paragraph("测试prompt")
#
#     # 将文档保存到内存中的 BytesIO 对象
#     buffer = BytesIO()
#     document.save(buffer)
#     buffer.seek(0)
#
#     export_btn.download_button(
#         "导出记录",
#         data=buffer,
#         file_name=f"{now:%Y-%m-%d %H.%M}_对录.docx",
#         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#         use_container_width=True,
#     )